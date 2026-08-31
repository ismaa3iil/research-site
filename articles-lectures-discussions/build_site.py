from __future__ import annotations

import html
import re
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).parent
SOURCE_DIR = ROOT / "sources"
ARTICLE_DIR = ROOT / "articles"

DOCUMENTS = [
    {
        "source": "Islam&Science-19.docx",
        "slug": "islam-and-science",
        "title": "الإسلام والعلم — أسئلة علمية كبرى من منظور إسلامي",
        "kind": "word",
    },
    {
        "source": "Journal Vol6 No2 ar.pdf",
        "slug": "arabic-journal-quality-education-volume-6-issue-2",
        "title": "المجلة العربية لجودة التعليم — المجلد ٦، العدد ٢",
        "kind": "pdf",
    },
    {
        "source": "الظاهرة المحمدية.docx",
        "slug": "the-muhammadan-phenomenon",
        "title": "الظاهرة المحمدية",
        "kind": "word",
    },
    {
        "source": "عن الأشعرية والسلفية.docx",
        "slug": "asharism-and-salafism",
        "title": "عن الأشعرية والسلفية",
        "kind": "word",
    },
    {
        "source": "مصداقيات.docx",
        "slug": "hadith-transmission-reliability",
        "title": "كيف نحسب مصداقيات الروايات المسنودة بحساب الاحتمالات",
        "kind": "word",
    },
    {
        "source": "مصطفى صبري واللانهاية.docx",
        "slug": "mustafa-sabri-and-infinity",
        "title": "مصطفى صبري واللانهاية",
        "kind": "word",
    },
    {
        "source": "ملاحظات سريعة على محاضرة الدكتور هشام غصيب3.docx",
        "slug": "notes-on-hisham-ghosheh-lecture",
        "title": "ملاحظات سريعة على محاضرة الدكتور هشام غصيب",
        "kind": "word",
    },
    {
        "source": "يحيى حمودة (2).doc",
        "slug": "yahia-hammoudeh",
        "title": "يحيى حمودة",
        "kind": "legacy",
    },
]


def esc(value: str) -> str:
    return html.escape(value, quote=True)


def natural_key(path: Path | str) -> tuple:
    name = Path(path).name
    return tuple(int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", name))


def document_paragraphs(path: Path) -> list[str]:
    document = Document(path)
    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.replace("\x00", "").strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def document_tables(path: Path) -> list[str]:
    document = Document(path)
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{esc(cell.text.strip())}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            tables.append("<table><tbody>" + "".join(rows) + "</tbody></table>")
    return tables


def extract_images(path: Path, article_path: Path) -> list[str]:
    image_dir = article_path / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    exported = []
    with ZipFile(path) as archive:
        media = sorted((name for name in archive.namelist() if name.startswith("word/media/")), key=natural_key)
        for name in media:
            suffix = Path(name).suffix.lower()
            if suffix not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                continue
            target = image_dir / Path(name).name
            target.write_bytes(archive.read(name))
            exported.append(target.name)
    return exported


def page_shell(title: str, body: str, *, description: str = "", nested: bool = False) -> str:
    asset_prefix = "../../" if nested else ""
    return f"""<!doctype html>
<html lang="ar" dir="rtl">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta name="description" content="{esc(description or title)}">
  <title>{esc(title)} — مقالات ومحاضرات ومناقشات</title>
  <link rel="stylesheet" href="{asset_prefix}styles.css">
</head>
<body>
  <main>{body}</main>
</body>
</html>
"""


def article_body(item: dict, paragraphs: list[str], tables: list[str], images: list[str]) -> str:
    root_prefix = "../../"
    if item["kind"] == "pdf":
        source_href = root_prefix + "sources/Journal%20Vol6%20No2%20ar.pdf"
        content = f"""
      <div class="notice">هذه المادة محفوظة بصيغة PDF الأصلية، ويمكن قراءتها مباشرة أو تنزيلها.</div>
      <object class="pdf-viewer" data="{source_href}" type="application/pdf">
        <p>لا يدعم المتصفح العرض المباشر. <a href="{source_href}">تنزيل الملف الأصلي</a>.</p>
      </object>
      <p class="source-link"><a class="button primary" href="{source_href}">تنزيل العدد بصيغة PDF</a></p>
"""
    elif item["kind"] == "legacy":
        source_href = root_prefix + "sources/" + esc(item["source"].replace(" ", "%20"))
        content = f"""
      <div class="notice">النسخة الأصلية محفوظة بصيغة Word القديمة (.doc). لم أقم بتحويلها تلقائيًا حتى لا يتغير نصها أو تنسيقها.</div>
      <p><a class="button primary" href="{source_href}">تنزيل الملف الأصلي</a></p>
"""
    else:
        paragraphs_html = "\n".join(f"      <p>{esc(text).replace(chr(10), '<br>')}</p>" for text in paragraphs)
        tables_html = "\n".join(f"      {table}" for table in tables)
        image_html = ""
        if images:
            image_items = "\n".join(
                f'          <a href="images/{esc(name)}"><img src="images/{esc(name)}" alt="صورة مرفقة من الملف الأصلي" loading="lazy"></a>'
                for name in images
            )
            image_html = f"\n      <h2>صور مرفقة</h2>\n      <div class=\"image-gallery\">\n{image_items}\n      </div>"
        source_href = root_prefix + "sources/" + esc(item["source"].replace(" ", "%20"))
        content = f"""
      <p class="source-link"><a class="button" href="{source_href}">تنزيل الملف الأصلي بصيغة Word</a></p>
      <div class="article-text">
{paragraphs_html}
{tables_html}
      </div>{image_html}
"""
    return f"""
    <a class="back-link" href="../../">العودة إلى الفهرس</a>
    <article class="article">
      <header class="article-header">
        <p class="eyebrow">قسم {DOCUMENTS.index(item) + 1} من ٨</p>
        <h1>{esc(item['title'])}</h1>
        <p class="intro-placeholder">ستضاف مقدمة موجزة لهذا النص لاحقًا.</p>
      </header>
      <div class="article-content">
{content}
      </div>
    </article>
"""


def build() -> None:
    SOURCE_DIR.mkdir(exist_ok=True)
    ARTICLE_DIR.mkdir(exist_ok=True)

    cards = []
    for item in DOCUMENTS:
        source = ROOT / item["source"]
        destination = SOURCE_DIR / item["source"]
        shutil.copy2(source, destination)
        article_path = ARTICLE_DIR / item["slug"]
        article_path.mkdir(parents=True, exist_ok=True)

        paragraphs: list[str] = []
        tables: list[str] = []
        images: list[str] = []
        extra = ""
        if item["kind"] == "word":
            paragraphs = document_paragraphs(source)
            tables = document_tables(source)
            images = extract_images(source, article_path)
            extra = f"{len(paragraphs)} فقرة"
            if tables:
                extra += f" · {len(tables)} جداول"
            if images:
                extra += f" · {len(images)} صور"
        elif item["kind"] == "pdf":
            pages = len(PdfReader(str(source)).pages)
            extra = f"{pages} صفحة · PDF"
        else:
            extra = "نسخة Word قديمة · DOC"

        article_html = page_shell(item["title"], article_body(item, paragraphs, tables, images), description=item["title"], nested=True)
        (article_path / "index.html").write_text(article_html, encoding="utf-8")
        cards.append(f"""
      <article class="paper">
        <p class="eyebrow">قسم {len(cards) + 1}</p>
        <h2><a href="articles/{item['slug']}/">{esc(item['title'])}</a></h2>
        <p class="meta">{esc(extra)}</p>
        <p>النص الأصلي محفوظ، وستضاف مقدمة تعريفية لهذا القسم لاحقًا.</p>
        <a class="button primary" href="articles/{item['slug']}/">قراءة القسم</a>
      </article>""")

    index_body = f"""
    <section class="hero">
      <p class="eyebrow">أرشيف إسماعيل حمودة</p>
      <h1>مقالات ومحاضرات ومناقشات</h1>
      <p class="lead">مجموعة من المقالات والمحاضرات والنصوص المحفوظة كما وردت في ملفاتها الأصلية.</p>
    </section>
    <section aria-labelledby="contents">
      <div class="section-heading">
        <h2 id="contents">الفهرس</h2>
        <span class="count">٨ أقسام</span>
      </div>
      <div class="papers">{''.join(cards)}
      </div>
    </section>
    <footer>تاريخ النشر الأولي: ٣١ آب ٢٠٢٦ · ستضاف المقدمات التعريفية لاحقًا.</footer>
"""
    (ROOT / "index.html").write_text(page_shell("مقالات ومحاضرات ومناقشات", index_body), encoding="utf-8")


if __name__ == "__main__":
    build()
